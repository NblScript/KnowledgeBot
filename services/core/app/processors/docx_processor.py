"""DOCX 文档处理器"""

import logging
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)


class DocxProcessor:
    """DOCX 文档处理器
    
    使用 python-docx 解析 Word 文档。
    """
    
    def __init__(self, extract_headers: bool = True, extract_tables: bool = True):
        """初始化 DOCX 处理器
        
        Args:
            extract_headers: 是否提取标题结构
            extract_tables: 是否提取表格
        """
        self.extract_headers = extract_headers
        self.extract_tables = extract_tables
    
    def parse(self, file_path: str) -> dict[str, Any]:
        """解析 DOCX 文件
        
        Args:
            file_path: DOCX 文件路径
            
        Returns:
            解析结果，包含：
            - text: 完整文本
            - paragraphs: 段落列表
            - tables: 表格列表
            - metadata: 文档元数据
        """
        try:
            doc = Document(file_path)
            
            paragraphs = []
            all_text = []
            current_section = {"title": "正文", "content": []}
            
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                
                # 判断是否为标题
                is_heading = self._is_heading(para)
                
                para_info = {
                    "text": text,
                    "style": para.style.name if para.style else None,
                    "is_heading": is_heading,
                    "level": self._get_heading_level(para) if is_heading else 0,
                }
                
                paragraphs.append(para_info)
                all_text.append(text)
            
            # 提取表格
            tables = []
            if self.extract_tables:
                for i, table in enumerate(doc.tables):
                    table_data = self._extract_table(table)
                    tables.append({
                        "index": i,
                        "data": table_data,
                        "text": self._table_to_text(table_data),
                    })
            
            # 提取元数据
            metadata = self._extract_metadata(doc)
            
            return {
                "text": "\n\n".join(all_text),
                "paragraphs": paragraphs,
                "tables": tables,
                "metadata": metadata,
                "total_paragraphs": len(paragraphs),
                "total_tables": len(tables),
            }
            
        except Exception as e:
            logger.error(f"Failed to parse DOCX {file_path}: {e}")
            raise
    
    def parse_with_structure(self, file_path: str) -> dict[str, Any]:
        """解析 DOCX 文件并保留文档结构
        
        Args:
            file_path: DOCX 文件路径
            
        Returns:
            解析结果，包含文档结构
        """
        try:
            doc = Document(file_path)
            
            sections = []
            current_section = None
            all_text = []
            
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                
                is_heading = self._is_heading(para)
                
                if is_heading:
                    # 保存上一个 section
                    if current_section:
                        sections.append(current_section)
                    
                    # 开始新 section
                    current_section = {
                        "title": text,
                        "level": self._get_heading_level(para),
                        "content": [],
                    }
                else:
                    if current_section is None:
                        current_section = {
                            "title": "文档开头",
                            "level": 0,
                            "content": [],
                        }
                    current_section["content"].append(text)
                
                all_text.append(text)
            
            # 添加最后一个 section
            if current_section:
                sections.append(current_section)
            
            # 提取表格
            tables = []
            if self.extract_tables:
                for i, table in enumerate(doc.tables):
                    table_data = self._extract_table(table)
                    tables.append({
                        "index": i,
                        "data": table_data,
                        "text": self._table_to_text(table_data),
                    })
            
            metadata = self._extract_metadata(doc)
            
            return {
                "text": "\n\n".join(all_text),
                "sections": sections,
                "tables": tables,
                "metadata": metadata,
            }
            
        except Exception as e:
            logger.error(f"Failed to parse DOCX with structure {file_path}: {e}")
            raise
    
    def _is_heading(self, para) -> bool:
        """判断段落是否为标题"""
        if para.style and para.style.name:
            return para.style.name.startswith("Heading") or "标题" in para.style.name
        return False
    
    def _get_heading_level(self, para) -> int:
        """获取标题级别"""
        if para.style and para.style.name:
            style_name = para.style.name
            if style_name.startswith("Heading"):
                try:
                    return int(style_name.replace("Heading", "").strip())
                except ValueError:
                    return 1
        return 1
    
    def _extract_table(self, table) -> list[list[str]]:
        """提取表格数据"""
        data = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                row_data.append(cell.text.strip())
            data.append(row_data)
        return data
    
    def _table_to_text(self, table_data: list[list[str]]) -> str:
        """将表格转换为文本"""
        lines = []
        for row in table_data:
            lines.append(" | ".join(row))
        return "\n".join(lines)
    
    def _extract_metadata(self, doc) -> dict[str, Any]:
        """提取文档元数据"""
        core_props = doc.core_properties
        
        return {
            "title": core_props.title or "",
            "author": core_props.author or "",
            "subject": core_props.subject or "",
            "keywords": core_props.keywords or "",
            "category": core_props.category or "",
            "comments": core_props.comments or "",
            "created": str(core_props.created) if core_props.created else "",
            "modified": str(core_props.modified) if core_props.modified else "",
            "last_modified_by": core_props.last_modified_by or "",
        }
