"""文档处理服务 - 文档解析、分块、向量化"""

import asyncio
import hashlib
import re
from abc import ABC, abstractmethod
from dataclasses import dataclass
from io import BytesIO
from typing import List

from app.config import settings


@dataclass
class ChunkData:
    """分块数据"""
    content: str
    chunk_index: int
    token_count: int
    metadata: dict = None
    
    def __post_init__(self):
        if self.metadata is None:
            self.metadata = {}


@dataclass
class ParsedDocument:
    """解析后的文档"""
    content: str
    file_name: str
    file_type: str
    metadata: dict = None
    
    def __post_init__(self):
        if self.metadata is None:
            self.metadata = {}


class DocumentParser(ABC):
    """文档解析器基类"""
    
    @abstractmethod
    async def parse(self, content: bytes, file_name: str) -> ParsedDocument:
        """解析文档内容"""
        pass
    
    @staticmethod
    def get_parser(file_type: str) -> "DocumentParser":
        """获取对应文件类型的解析器"""
        parsers = {
            "pdf": PDFParser,
            "txt": TextParser,
            "md": MarkdownParser,
            "docx": DocxParser,
            "html": HTMLParser,
        }
        parser_class = parsers.get(file_type.lower())
        if not parser_class:
            raise ValueError(f"Unsupported file type: {file_type}")
        return parser_class()


class TextParser(DocumentParser):
    """纯文本解析器"""
    
    async def parse(self, content: bytes, file_name: str) -> ParsedDocument:
        text = content.decode("utf-8", errors="ignore")
        return ParsedDocument(
            content=text,
            file_name=file_name,
            file_type="txt",
            metadata={"encoding": "utf-8"},
        )


class MarkdownParser(DocumentParser):
    """Markdown 解析器"""
    
    async def parse(self, content: bytes, file_name: str) -> ParsedDocument:
        text = content.decode("utf-8", errors="ignore")
        
        # 提取标题作为元数据
        headers = re.findall(r"^#+\s+(.+)$", text, re.MULTILINE)
        
        return ParsedDocument(
            content=text,
            file_name=file_name,
            file_type="md",
            metadata={
                "headers": headers[:10],  # 最多保留10个标题
                "encoding": "utf-8",
            },
        )


class PDFParser(DocumentParser):
    """PDF 解析器"""
    
    async def parse(self, content: bytes, file_name: str) -> ParsedDocument:
        try:
            from pypdf import PdfReader
        except ImportError:
            raise ImportError("pypdf not installed. Run: pip install pypdf")
        
        # 在线程池中运行，避免阻塞
        loop = asyncio.get_event_loop()
        result = await loop.run_in_executor(
            None,
            self._parse_pdf,
            content,
            file_name,
        )
        return result
    
    def _parse_pdf(self, content: bytes, file_name: str) -> ParsedDocument:
        reader = PdfReader(BytesIO(content))
        text_parts = []
        metadata = {"pages": len(reader.pages)}
        
        for i, page in enumerate(reader.pages):
            page_text = page.extract_text() or ""
            if page_text.strip():
                text_parts.append(f"[Page {i+1}]\n{page_text}")
        
        full_text = "\n\n".join(text_parts)
        
        return ParsedDocument(
            content=full_text,
            file_name=file_name,
            file_type="pdf",
            metadata=metadata,
        )


class DocxParser(DocumentParser):
    """Word DOCX 解析器"""
    
    async def parse(self, content: bytes, file_name: str) -> ParsedDocument:
        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx not installed. Run: pip install python-docx")
        
        loop = asyncio.get_event_loop()
        result = await loop.run_in_executor(
            None,
            self._parse_docx,
            content,
            file_name,
        )
        return result
    
    def _parse_docx(self, content: bytes, file_name: str) -> ParsedDocument:
        doc = Document(BytesIO(content))
        
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)
        
        # 提取表格内容
        tables_text = []
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text for cell in row.cells)
                tables_text.append(row_text)
        
        full_text = "\n\n".join(paragraphs)
        if tables_text:
            full_text += "\n\n[Tables]\n" + "\n".join(tables_text)
        
        return ParsedDocument(
            content=full_text,
            file_name=file_name,
            file_type="docx",
            metadata={"paragraphs": len(paragraphs), "tables": len(doc.tables)},
        )


class HTMLParser(DocumentParser):
    """HTML 解析器"""
    
    async def parse(self, content: bytes, file_name: str) -> ParsedDocument:
        try:
            from bs4 import BeautifulSoup
        except ImportError:
            # 如果没有 BeautifulSoup，简单去除 HTML 标签
            text = content.decode("utf-8", errors="ignore")
            text = re.sub(r"<[^>]+>", "", text)
            return ParsedDocument(
                content=text,
                file_name=file_name,
                file_type="html",
            )
        
        loop = asyncio.get_event_loop()
        result = await loop.run_in_executor(
            None,
            self._parse_html,
            content,
            file_name,
        )
        return result
    
    def _parse_html(self, content: bytes, file_name: str) -> ParsedDocument:
        from bs4 import BeautifulSoup
        
        soup = BeautifulSoup(content, "html.parser")
        
        # 移除脚本和样式
        for element in soup(["script", "style", "nav", "footer"]):
            element.decompose()
        
        text = soup.get_text(separator="\n")
        # 清理多余空白
        lines = [line.strip() for line in text.splitlines() if line.strip()]
        clean_text = "\n".join(lines)
        
        # 提取标题
        title = soup.find("title")
        title_text = title.get_text() if title else file_name
        
        return ParsedDocument(
            content=clean_text,
            file_name=file_name,
            file_type="html",
            metadata={"title": title_text},
        )


class TextChunker:
    """文本分块器"""
    
    def __init__(
        self,
        chunk_size: int = None,
        chunk_overlap: int = None,
    ):
        self.chunk_size = chunk_size or settings.chunk_size
        self.chunk_overlap = chunk_overlap or settings.chunk_overlap
    
    def chunk(self, text: str, metadata: dict = None) -> List[ChunkData]:
        """将文本分割成块"""
        if not text.strip():
            return []
        
        # 按段落分割
        paragraphs = text.split("\n\n")
        
        chunks = []
        current_chunk = ""
        chunk_index = 0
        
        for para in paragraphs:
            para = para.strip()
            if not para:
                continue
            
            # 如果单个段落超过 chunk_size，需要进一步分割
            if len(para) > self.chunk_size:
                # 先处理当前累积的内容
                if current_chunk:
                    chunks.append(self._create_chunk(current_chunk, chunk_index, metadata))
                    chunk_index += 1
                    current_chunk = ""
                
                # 按句子分割长段落
                for sub_chunk in self._split_long_text(para):
                    chunks.append(self._create_chunk(sub_chunk, chunk_index, metadata))
                    chunk_index += 1
            else:
                # 检查是否需要创建新块
                if len(current_chunk) + len(para) + 2 > self.chunk_size:
                    if current_chunk:
                        chunks.append(self._create_chunk(current_chunk, chunk_index, metadata))
                        chunk_index += 1
                    current_chunk = para
                else:
                    if current_chunk:
                        current_chunk += "\n\n" + para
                    else:
                        current_chunk = para
        
        # 处理最后一块
        if current_chunk:
            chunks.append(self._create_chunk(current_chunk, chunk_index, metadata))
        
        return chunks
    
    def _split_long_text(self, text: str) -> List[str]:
        """分割长文本"""
        # 按句子分割
        sentences = re.split(r"[。！？.!?]", text)
        sentences = [s.strip() for s in sentences if s.strip()]
        
        chunks = []
        current = ""
        
        for sent in sentences:
            if len(current) + len(sent) + 1 > self.chunk_size:
                if current:
                    chunks.append(current)
                current = sent
            else:
                if current:
                    current += "。" + sent
                else:
                    current = sent
        
        if current:
            chunks.append(current)
        
        return chunks
    
    def _create_chunk(self, content: str, index: int, base_metadata: dict = None) -> ChunkData:
        """创建分块数据"""
        # 简单估算 token 数（中文约 1.5 字/token，英文约 4 字/token）
        chinese_chars = len(re.findall(r"[\u4e00-\u9fff]", content))
        other_chars = len(content) - chinese_chars
        token_count = int(chinese_chars / 1.5 + other_chars / 4)
        
        return ChunkData(
            content=content,
            chunk_index=index,
            token_count=token_count,
            metadata=base_metadata or {},
        )


class DocumentProcessor:
    """文档处理器 - 组合解析、分块、向量化"""
    
    def __init__(self):
        self.chunker = TextChunker()
    
    async def process(
        self,
        content: bytes,
        file_name: str,
        file_type: str,
    ) -> tuple[ParsedDocument, List[ChunkData]]:
        """
        处理文档：解析 -> 分块
        
        Returns:
            (解析后的文档, 分块列表)
        """
        # 1. 解析文档
        parser = DocumentParser.get_parser(file_type)
        parsed = await parser.parse(content, file_name)
        
        # 2. 分块
        chunks = self.chunker.chunk(parsed.content, parsed.metadata)
        
        return parsed, chunks
    
    async def embed_chunks(
        self,
        chunks: List[ChunkData],
        kb_id: str,
        doc_id: str,
    ) -> List[dict]:
        """
        将分块向量化
        
        Returns:
            向量化后的数据，可直接写入 Milvus
        """
        from app.services.embedding import get_embedding
        
        embedding_service = get_embedding()
        
        # 批量处理，每次最多 100 个
        batch_size = 100
        all_vectors = []
        
        for i in range(0, len(chunks), batch_size):
            batch = chunks[i:i + batch_size]
            texts = [c.content for c in batch]
            
            result = await embedding_service.embed(texts)
            
            for j, (chunk, vector) in enumerate(zip(batch, result.vectors)):
                chunk_id = hashlib.md5(f"{doc_id}_{chunk.chunk_index}".encode()).hexdigest()[:16]
                all_vectors.append({
                    "id": chunk_id,
                    "vector": vector,
                    "content": chunk.content,
                    "doc_id": doc_id,
                    "kb_id": kb_id,
                    "chunk_index": chunk.chunk_index,
                    "token_count": chunk.token_count,
                    "metadata": chunk.metadata,
                })
        
        return all_vectors


# 全局实例
_processor: DocumentProcessor | None = None


def get_processor() -> DocumentProcessor:
    """获取文档处理器"""
    global _processor
    if _processor is None:
        _processor = DocumentProcessor()
    return _processor
